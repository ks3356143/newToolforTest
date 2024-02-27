from PyQt5 import QtCore
from PyQt5.QtCore import pyqtSignal
from pathlib import *
from PyQt5.QtWidgets import QMessageBox
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import io

class create_new_JtoS(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self, parent):
        super().__init__()
        self.parent = parent

    def run(self):
        self.sin_out.emit("进入CPU测试记录转说明线程......")
        self.sin_out.emit("开始转换......")
        # 如果没有选择文件路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return

        # 打开模板文件进行渲染，然后就是用docxtpl生成用例
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "cpu新记录to说明模版.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path)  # 模板导入成功

        except:
            QMessageBox.warning(self.parent, "出错了", "导入模板出错请检查模板文件是否存在或名字不正确")
            return

        try:
            doc = Document(self.parent.open_file_name[0])
            self.sin_out.emit('已识别到CPU测试记录文件...')
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.parent.tabWidget.setEnabled(True)
            return

        self.sin_out.emit('复制测试说明文档模板到本程序所在目录...')
        curpath = Path.cwd() / 'need'
        shuoming_path_tmp = curpath / 'document_templates' / 'cpu新记录to说明模版.docx'
        print(shuoming_path_tmp)
        if shuoming_path_tmp.is_file():
            self.sin_out.emit('已检测到有说明模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return

        # 获取表格数量
        tables = doc.tables
        tb_count = len(tables)
        self.sin_out.emit('total:' + str(tb_count))
        # 创建一个字典来储存单个用例
        data_list = []
        table_index = 1
        # 获取表格数量
        for ele in doc._element.body.iter():
            data = {'type': ''}
            if ele.tag.endswith('}p'):
                elePstyle = ele.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if len(elePstyle) >= 1:
                    parag = Paragraph(ele, doc)
                    if parag.style.name.startswith("Heading") or parag.style.name.startswith("标题"):
                        rank = parag.style.name.split(" ")[-1]
                        text = parag.text
                        if text == "" or text.startswith("文档齐套性审查单") \
                                    or text.startswith("软件研制任务书审查单") \
                                    or text.startswith("附录") \
                                    or text.startswith("附件") \
                                    or text.startswith("需求文档审查单"):
                            pass
                        else:
                            data['type'] = rank
                            data['title'] = text
                            data_list.append(data)
            elif ele.tag.endswith('}tbl'):
                data = {'name': '', 'biaoshi': '', 'zhuizong': [], 'zongsu': '', 'init': '', 'qianti': '', 'step': []}
                data['type'] = 'table'
                table = Table(ele, doc)
                if table.cell(1, 0).text == '测试用例名称':
                    self.sin_out.emit(str(table_index))

                    try:
                        self.sin_out.emit(str(table_index))
                        self.sin_out.emit(f'正在处理第{table_index}个表格')
                        # 1、获取测试用例名称
                        data['name'] = table.cell(1, 3).text
                        # 2、获取用例标识
                        data['biaoshi'] = table.cell(1, 9).text
                        # 3、获取追踪关系 注意word中换行为\r\x07
                        temp = table.cell(2, 3)
                        for tem in temp.paragraphs:
                            data['zhuizong'].append(tem.text)
                        # 4、获取综述
                        data['zongsu'] = table.cell(3, 3).text
                        # 5、初始化
                        data['init'] = table.cell(4, 3).text
                        # 6、获取前提与约束
                        data['qianti'] = table.cell(5, 3).text
                        # 7、获取步骤信息-总行数减去12为步骤行数
                        row_count = len(table.rows)
                        step_count = row_count - 12
                        for j in range(step_count):
                            buzhou = {'shuru': '', 'yuqi': '', 'num': '', 'image': '', 'is_image': '0'}
                            buzhou['num'] = table.rows[8 + j].cells[0].text
                            if buzhou['num'] == '':
                                buzhou['num'] = j + 1
                            buzhou['shuru'] = table.rows[8 + j].cells[2].text
                            cel = table.rows[8 + j].cells[2]
                            if len(
                                cel._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            ) > 0:
                                img_ele = cel._element.xpath('.//pic:pic')[0]
                                embed = img_ele.xpath('.//a:blip/@r:embed')[0]
                                related_part = doc.part.related_parts[embed]
                                image = related_part.image
                                # blob属性就是二进制图片属性
                                image_bytes = image.blob
                                buzhou['image'] = InlineImage(tpl, io.BytesIO(image_bytes))
                                buzhou['is_image'] = '1'
                            buzhou['yuqi'] = table.rows[8 + j].cells[4].text
                            data['step'].append(buzhou)
                        # 8、最后加入data_list
                        data_list.append(data)
                        table_index += 1
                    except:
                        self.sin_out.emit(f'第{table_index}个表格处理错误！')
                        table_index += 1
                        pass
        # 开始渲染模板文件
        try:
            self.sin_out.emit('all_doned:')
            context = {
                "tables": data_list,
                "renyuan": self.parent.lineEdit_17.text(),
            }
            tpl.render(context)
            tpl.save("CPU新版反向生成说明.docx")
            self.sin_out.emit('stopsuccess')
        except:
            self.sin_out.emit('stoperror')
            return
