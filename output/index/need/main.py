# -*- coding: utf-8 -*-
import logging
LOG_FORMAT = "%(asctime)s>%(levelname)s>PID:%(process)d %(thread)d>%(module)s>%(funcName)s>%(lineno)d>%(message)s"
logging.basicConfig(level=logging.DEBUG, format=LOG_FORMAT, )

# 是否打印调试信息标志
debug = True
if debug:
    logging.debug("进入主程序，开始导入包...")
    
#导入常规库
import sys,re,string 
from pathlib import *
#导入word文档操作库
from win32com.client import DispatchEx
from docxtpl import DocxTemplate
import docx
from docx import Document
import shutil
import pythoncom
#导入QT组件
from PyQt5 import QtCore
from PyQt5.QtWidgets import QMainWindow,QFileDialog,QMessageBox,QDialog,QToolTip
from PyQt5.QtCore import pyqtSignal
from PyQt5.QtGui import QFont
#导入UI转换PY文件
from need.Ui_GUI import Ui_MainWindow
from need import about,zhuan
#导入工具包文件-时间转换
from need.utils import get_current_time,get_current_name,get_current_date,get_current_hour
from need.zhuan_tool import IEEE754_16_to_float,IEEE754_float_to_16

class zhuan_dlg(QDialog,zhuan.Ui_Dialog):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        
        if debug:
            logging.debug("初始化转换程序:")
            
        #linetext信号连接
        self.lineEdit.editingFinished.connect(self.shiliu_zhuan)
        self.lineEdit_2.editingFinished.connect(self.shi_zhuan)
        # 设置气泡提示信息
        QToolTip.setFont(QFont("SansSerif", 12))
        self.lineEdit.setToolTip("注意，编辑输入框后需要点击其他控件才进行转换，且格式不正确不转换")
        self.lineEdit_2.setToolTip("注意，编辑输入框后需要点击其他控件才进行转换，且格式不正确不转换")
        
    def shiliu_zhuan(self):
        #获取当前文字
        x = self.lineEdit.text()
        if len(x) == 8:
            if self.radioButton.isChecked() == True: #说明选中了32位转换
                result = IEEE754_16_to_float(x,32)
                self.lineEdit_2.setText(str(result))
        elif len(x) == 16:
            if self.radioButton_2.isChecked() == True: #说明选中了64位转换
                result = IEEE754_16_to_float(x,64)
                self.lineEdit_2.setText(str(result))
        
    def shi_zhuan(self):
        #获取当前文字
        x = self.lineEdit_2.text()
        #字符串转浮点数
        if str.isdigit(x):
            x = float(x)
        elif str.isdigit(x.replace(".","")):
            x = float(x)
        if isinstance(x,float) == True:
            if self.radioButton.isChecked() == True: #说明选中了32位转换
                result = IEEE754_float_to_16(x,32)
                self.lineEdit.setText(str(result))
            else:
                result = IEEE754_float_to_16(x,64)
                self.lineEdit.setText(str(result))
             
class userMain(QMainWindow,Ui_MainWindow):
    #自定义信号和槽
    
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        if debug:
            logging.debug("初始化主程序:")
        
        # 实例化翻译家
        self.trans = QtCore.QCoreApplication.translate
        self.setWindowTitle('测试个人工具')
        
        #使用翻译家改变PYQT的空间名字等属性self.label_4.setText(self.trans("MainWindow", "文件名（自动识别）："))
        
        if debug:
            logging.debug("初始化部分全局变量...")
        #存放文件夹路径变量
        self.open_dirs_name = ''
        #存放文件名称路径变量
        self.open_file_name = ''
        
        #读取配置文件
        
        #~~~~~~~~~~~~~~~连接线程函数~~~~~~~~~~~~~~~
        ##连接大纲生成说明函数
        self.create_shuoming_trd = create_shuoming(self) 
        self.create_shuoming_trd.sin_out.connect(self.text_display) 
        self.pushButton_2.clicked.connect(self.create_shuoming_btn) 
        
        ##连接大纲追溯
        self.create_dagang_zhuisu_trd = create_dagang_zhuisu(self) 
        self.create_dagang_zhuisu_trd.sin_out.connect(self.text_display) 
        self.pushButton_5.clicked.connect(self.creat_shuoming_zhuisu_btn) 

        ##连接说明追踪线程
        self.create_shuoming_zhuisu_trd = create_shuoming_zhuisu(self)  
        self.create_shuoming_zhuisu_trd.sin_out.connect(self.text_display)  
        self.pushButton_6.clicked.connect(self.creat_dagang_zhuisu_btn)  

        ##连接报告追踪
        self.create_baogao_zhuisu_trd = create_baogao_zhuisu(self)  
        self.create_baogao_zhuisu_trd.sin_out.connect(self.text_display)  
        self.pushButton_18.clicked.connect(self.create_baogao_zhuisu_btn) 
        
        ##连接单元追踪线程
        self.create_danyuan_trd = create_danyuan(self) 
        self.create_danyuan_trd.sin_out.connect(self.text_display) 
        self.pushButton_8.clicked.connect(self.creat_danyuan_btn) 
        
        ##连接根据测试说明生成记录线程
        self.create_jilu_trd = create_jilu(self) 
        self.create_jilu_trd.sin_out.connect(self.text_display) 
        self.pushButton_12.clicked.connect(self.creat_jilu_btn) 
        
        ##记录反向生成说明线程
        self.create_shuomingfanxiang_trd = create_shuomingfanxiang(self) 
        self.create_shuomingfanxiang_trd.sin_out.connect(self.text_display) 
        self.pushButton_13.clicked.connect(self.creat_shuomingfanxiang_btn) 
        
        ##自动填充空白表格线程
        self.create_zidong_trd = create_zidong(self) 
        self.create_zidong_trd.sin_out.connect(self.text_display) 
        self.pushButton_15.clicked.connect(self.creat_zidong_btn)
        
        ##清空单元格线程
        self.clear_cell_trd = clear_cell(self) 
        self.clear_cell_trd.sin_out.connect(self.text_display) 
        self.pushButton_14.clicked.connect(self.clear_cell_btn)
        
        ##提取表格内容线程
        self.get_content_trd = get_content(self) 
        self.get_content_trd.sin_out.connect(self.text_display) 
        self.pushButton_19.clicked.connect(self.get_content_btn)
        
        #自定义信号连接
        
        # 获取状态栏对象
        self.user_statusbar = self.statusBar()
        # 右下角窗口尺寸调整符号
        self.user_statusbar.setSizeGripEnabled(True)
        self.user_statusbar.setStyleSheet("QStatusBar.item{border:10px}")
        
        #~~~~~~~~~~~~~~~按钮连接函数~~~~~~~~~~~~~~~~
        ##选择文件按钮连接
        self.pushButton.clicked.connect(self.choose_docx_func)
        self.pushButton_4.clicked.connect(self.choose_docx_func)
        self.pushButton_7.clicked.connect(self.choose_docx_func)
        self.pushButton_11.clicked.connect(self.choose_docx_func)
        self.pushButton_16.clicked.connect(self.choose_docx_func)
        self.pushButton_17.clicked.connect(self.choose_docx_func)
        #清空显示区
        self.pushButton_9.clicked.connect(self.clear_textEdit_content)
        #显示帮助
        self.pushButton_10.clicked.connect(self.display_help)
        
        #~~~~~~~~~~~~~~~导航栏按钮连接函数~~~~~~~~~~~~~~~~
        #显示关于本软件的菜单
        self.actionAbout.triggered.connect(self.display_about)
        #打开文件夹
        self.actionopen.triggered.connect(self.choose_docx_func)
        #打开IEEE754转换工具
        self.actionIEEE754.triggered.connect(self.open_zhuan_tool)
        
        if debug:
            logging.debug("界面加载完成...")
        
#~~~~~~~~~~~~~~~~~~~~初始化直接运行的函数（也就是起始运行一次）~~~~~~~~~~~~~~~~~~~~~~~~~~


#~~~~~~~~~~~~~~~~~~~~间接按钮函数，用户点击后操作~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    #清空显示区函数
    def clear_textEdit_content(self):
        self.textBrowser.clear()
        return
    #显示帮助函数
    def display_help(self):
        txt_path = Path.cwd() / 'need' /'others' /'readme.txt'
        with open(txt_path, 'r',encoding='utf-8') as f:
            data = f.read()
            self.textBrowser.append(data)
        return
    #显示关于函数
    def display_about(self):
        dlg = QDialog()
        about_dlg = about.Ui_Dialog()
        about_dlg.setupUi(dlg)
        dlg.show()
        dlg.exec_()
        print("显示关于界面")
        return
    
    #显示IEEE754转换工具
    def open_zhuan_tool(self):
        dlg_zhuan = zhuan_dlg() #实例化界面
        dlg_zhuan.show()
        
        
        dlg_zhuan.exec_()
        
        return


#~~~~~~~~~~~~~~~~~~~~线程区域函数，用于启动线程~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#生成测试说明启动函数
    def create_shuoming_btn(self):
        self.create_shuoming_trd.start()
        self.tabWidget.setEnabled(False)
        return
    
# 大纲追溯线程启动函数
    def creat_dagang_zhuisu_btn(self):
        self.create_dagang_zhuisu_trd.start()
        self.tabWidget.setEnabled(False)

# 说明追溯线程启动函数
    def creat_shuoming_zhuisu_btn(self):
        self.create_shuoming_zhuisu_trd.start()
        self.tabWidget.setEnabled(False)

# 提取单元格标题右侧内容线程启动函数
    def create_baogao_zhuisu_btn(self):
        self.create_baogao_zhuisu_trd.start()
        self.tabWidget.setEnabled(False)

# 记录反向生成说明线程
    def creat_shuomingfanxiang_btn(self):
        self.create_shuomingfanxiang_trd.start()
        self.tabWidget.setEnabled(False)
        
# 单元测试报告转换为我们的用例线程
    def creat_danyuan_btn(self):
        self.create_danyuan_trd.start()
        self.tabWidget.setEnabled(False)
        
# 单元测试报告转换为我们的用例线程
    def creat_jilu_btn(self):
        self.create_jilu_trd.start()
        self.tabWidget.setEnabled(False)
        
# 自动填充线程
    def creat_zidong_btn(self):
        self.create_zidong_trd.start()
        self.tabWidget.setEnabled(False)
        
# 清空表格单元格内容
    def clear_cell_btn(self):
        self.clear_cell_trd.start()
        self.tabWidget.setEnabled(False)

# 提取单元格标题右侧内容线程启动函数
    def get_content_btn(self):
        self.get_content_trd.start()
        self.tabWidget.setEnabled(False)
    
#选择文档函数
    def choose_docx_func(self):
        self.open_file_name = QFileDialog.getOpenFileName(
            self, '选择文件', '.', "Word files(*.docx)")
        self.textBrowser.append('已选择文件路径：' + self.open_file_name[0])
        
#关闭线程函数
    def stop_shuoming_thread(self):
        self.tabWidget.setEnabled(True)
        self.create_shuoming_trd.terminate()
        self.create_dagang_zhuisu_trd.terminate()
        self.create_shuoming_zhuisu_trd.terminate()
        self.create_shuomingfanxiang_trd.terminate()
        self.create_jilu_trd.terminate()
        self.create_zidong_trd.terminate()
        self.clear_cell_trd.terminate()
        self.get_content_trd.terminate()
        print("停止线程成功！")
    
#~~~~~~~~~~~~~~~~~~~~显示函数~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    def text_display(self, texttmp):
        if texttmp[:10] == 'stopthread':
            self.stop_shuoming_thread()
            return
    
        if texttmp[:6] == 'total:':
            cnt = int(texttmp[6:])
            self.progressBar.setRange(0, cnt - 1)

        if texttmp == 'function success':
            QMessageBox.information(self, '操作成功', '请查看本程序当前文件夹下的相关文档！')
            self.textBrowser.append('完成！！！')
            self.tabWidget.setEnabled(True)
            return
        if texttmp == '保存文件错':
            QMessageBox.warning(self, '出错了', '保存文件失败！')
            self.tabWidget.setEnabled(True)
            return
        if texttmp == 'no folder':
            QMessageBox.information(self, '没有选择文件夹', '还没有选择文件夹，点击"文件"菜单进行选择！')
            return
        if texttmp.find('warning:') != -1:
            QMessageBox.information(self, 'WARNING', texttmp[8:])
            return

        if texttmp.find('open failed:') != -1:
            QMessageBox.warning(
                self, '打开文件失败',
                '打开' + texttmp[12:] + '失败' + '请确认文档是否打开或者模板文件存在且后缀名为docx！')
            return
        if texttmp == 'nofile':
            QMessageBox.information(self, '错误',
                                    '还没有选择文件（夹），点击"文件"菜单或者工具栏进行选择！')
            return
        if texttmp.isdigit() == True:
            self.progressBar.setValue(int(texttmp))
        else:
            self.textBrowser.append(texttmp)

    def closeEvent(self, event):
        reply = QMessageBox.question(self, '提示',
                    "是否要关闭所有窗口?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No)
        if reply == QMessageBox.Yes:
            event.accept()
            sys.exit(0)   # 退出程序
        else:
            event.ignore()
    
##################################################################################
#大纲生成测试说明线程
##################################################################################
class create_shuoming(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        #用来储存测试项DC等转换
        zhuan_dict = {'DC':'文档审查','SU':'功能测试','CR':'代码审查','SA':'静态分析','AC':'性能测试',\
            'IO':'接口测试','SE':'安全性测试','BT':'边界测试','RE':'恢复性测试','ST':'强度测试',\
                'AT':'余量测试','GUI':'人机交互界面测试','DP':'数据处理测试','JR':'兼容性测试',\
                    'LG':'逻辑测试','AZ':'安装性测试','TT':'时序测试','PA':'功耗分析'}
        
        self.sin_out.emit("进入军品大纲转说明......")
        self.sin_out.emit("开始转换......")
        #如果没有选择路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        #告诉windows单线程
        pythoncom.CoInitialize()
        #在用户选择的目录中查找大纲文档
        self.sin_out.emit('打开测评大纲文档...')
        
        #使用win32com打开-记得关闭
        #打开word应用
        self.w = DispatchEx('Word.Application')
        #self.w.visible=True
        self.w.DisplayAlerts = 0
        try:
            dagangfile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        self.sin_out.emit('复制测试说明文档模板到本程序所在目录...')
        curpath = Path.cwd() / 'need'
        shuoming_path_tmp = curpath / 'document_templates' / '测试说明模板.docx'
        print(shuoming_path_tmp)
        if shuoming_path_tmp.is_file():
            self.sin_out.emit('已检测到有说明模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return
        
        #创建一个字典来储存单个用例
        
        data_list = []
        
        #获取表格数量
        try:
            csx_tb_count = dagangfile.Tables.Count
            self.sin_out.emit('total:'+ str(csx_tb_count))
        except:
            self.sin_out.emit('不存在表格！')
            QMessageBox.warning(self.parent,'出错了','测试说明文档格式错误或者没有正确表格')
            try:
                dagangfile.Close()
            except:
                pass
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        #循环表格
        yongli_count = 0
        #用来储存章节号中的DC、SU等标识，用于章节号判断
        is_fire_su = ""
        #用来储存基本_分割后个数
        num_fenge = 3
        for i in range(csx_tb_count):
            self.sin_out.emit(str(i))
            if dagangfile.Tables[i].Rows.Count > 2:
                #注意win32com的Cell从1开始不是从0开始
                if dagangfile.Tables[i].Cell(1, 1).Range.Text.find('测试项名称') != -1:               
                    #一个用例不变内容获取
                    dagangfile.Tables[i].Rows.First.Select()
                    zhangjiehao = self.w.Selection.Bookmarks("\headinglevel").\
                        Range.Paragraphs(1).Range.ListFormat.ListString
                    zhangjieming = self.w.Selection.Bookmarks("\headinglevel").\
                        Range.Paragraphs(1).Range.Text.rstrip('\r')
                    print("测试项所在章节号：",zhangjiehao)
                    #获取用例标识不加上序号_1
                    basic_biaoshi = dagangfile.Tables[i].Cell(1,4).Range.Text.rstrip()[:-2]
                    print("测试项标识为：",basic_biaoshi)
                    
                    #储存num_fenge的数值，初始化
                    if yongli_count == 0:
                        num_fenge = len(basic_biaoshi.split("_"))
                    #获取测试用例名称Cell(4,2)整行
                    info_ceshi_buzhou = dagangfile.Tables[i].Cell(4,2).Range.Text
                    info_ceshi_yuqi = dagangfile.Tables[i].Cell(9,2).Range.Text
                    #判断是否只有一行，如果只有一行处理表格
                    if dagangfile.Tables[i].Cell(4, 2).Range.Paragraphs.Count <= 1:
                        
                        #缓存一个用例的data填入数据
                        data = {'zhangjie':'','mingcheng':'','biaoshi':'','is_first':'1', \
                        'yueshu':'软件正常工作，环境连接正常', 'yongli_biaoshi':'','renyuan':'陈俊亦',\
                        'chushi':'外接设备或软件运行正常','csx_mingcheng':'','is_begin':'0',\
                            'zongsu':'',"zuhe":[],'csxbs':""}
                        zuhe_dict = {"buzhou":"","yuqi":"","xuhao":"1"}
                        try:
                            #填写一行情况下表格
                            data['mingcheng'] = dagangfile.Tables[i].Cell(1,2).Range.Text.rstrip('\r\x07')
                            #注意word中后面都有2个字符
                            data['yongli_biaoshi'] = (basic_biaoshi + "_1").replace('XQ','YL')
                            data['zhangjie'] = zhangjiehao
                            data['csx_mingcheng'] = zhangjieming
                            data['biaoshi'] = basic_biaoshi
                            data['zongsu'] = dagangfile.Tables[i].Cell(3,2).Range.Text[:-2]
                            
                            zuhe_dict["buzhou"] = dagangfile.Tables[i].Cell(4,2).Range.Text.rstrip('\r\x07')
                            zuhe_dict["yuqi"] = dagangfile.Tables[i].Cell(9,2).Range.Text.rstrip('\r\x07')
                            zuhe_dict["xuhao"] = '1'
                            data["zuhe"].append(zuhe_dict)
                            #判断是否为第一个测试类型，如果是修改章节号标识,则将章节号展示，如果和储存相同
                            #则章节号不展示
                            #首先获取测试项标识，分割成列表
                            fenge = data['biaoshi'].split("_")
                            #获取当前测试项分割后的个数
                            if len(fenge) < num_fenge:
                                if fenge[-1]!= is_fire_su:
                                    is_fire_su = fenge[-1]
                                    data['is_begin'] = "1"
                                    data['csxbs'] = zhuan_dict[fenge[-1]]
                            else:
                                if fenge[-2] != is_fire_su:
                                    is_fire_su = fenge[-2]
                                    data['is_begin'] = "1"
                                    data['csxbs'] = zhuan_dict[fenge[-2]]
                            if self.parent.lineEdit.text():
                                data['renyuan'] = self.parent.lineEdit.text()
                            #将data加入data_list
                            data['is_first'] = "1"
                            data_list.append(data)
                            yongli_count += 1
                            
                            
                            self.sin_out.emit('###获取用例序号：{}'.format(yongli_count))
                            self.sin_out.emit('###该用例标识为：{}'.format(data['yongli_biaoshi']))
                        except:
                            self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格处理失败$$$$$$$$$$$$')
                            pass
                        
                    elif dagangfile.Tables[i].Cell(4, 2).Range.Paragraphs.Count > 1:
                        
                        try:
                            #下面拆分每行，使用\r（回车）分割
                            info_buzhou_list = list(filter(lambda x:x!="\x07" and x!="",info_ceshi_buzhou.split('\r')))
                            info_yuqi_list = list(filter(lambda x:x!="\x07" and x!="",info_ceshi_yuqi.split('\r')))
                            
                            #去掉括号和以下字符
                            rule = "[()（）;；。]" #rule为去掉的符号（这个可以改TODO）

                            #初始化去掉rule的列表
                            buzhou_list = []
                            yuqi_list = []
                            
                            for item in info_buzhou_list:
                                buzhou_list.append(re.sub(rule,"",item).strip(string.digits))
                            for item in info_yuqi_list:
                                yuqi_list.append(re.sub(rule,"",item).strip(string.digits))
                            
                            #获取测试项综述-为该循环前不变内容
                            basic_zongshu = buzhou_list.pop(0).strip()
                            print('获取的测试用例综述是：',data["zongsu"])
                            
                            #获取字典中的buzhou和yuqi,找冒号
                            j = -1 #自制列表索引
                            substrict_list = [] #差值列表
                            for item in buzhou_list:
                                #先找到冒号所在索引
                                j = j + 1
                                if item.find(":") != -1 or item.find("：") != -1:
                                    #现在知道冒号所在行号，现在要确定每个用例几行
                                    substrict_list.append(j)
                                    
                            #！！！注意差值计算步骤需要-1才是正确的步骤数量
                            # self.sin_out.emit("解析测试项序号"+ str(i) + "|检测到冒号所在行号为：" \
                            #     + str(substrict_list))
                            # self.sin_out.emit("|检测到步骤总行数(序号)：" \
                            #     + str(len(buzhou_list)))
                            
                            #循环用例个数
                            count_test = len(substrict_list)
                            temp_list = substrict_list
                            temp_list.append(len(buzhou_list))
                            
                            for item in range(count_test):
                                #初始化data数据,包括步骤和预期、序号dict
                                data = {'zhangjie':'','mingcheng':'','biaoshi':'','is_first':'0', \
                                    'yueshu':'软件正常工作，环境连接正常', 'yongli_biaoshi':'','renyuan':'陈俊亦',\
                                    'chushi':'外接设备或软件运行正常','csx_mingcheng':'','is_begin':'0',\
                                        'zongsu':'',"zuhe":[],'csxbs':""}
                                
                                #这里要求冒号最后一个
                                data['mingcheng'] = buzhou_list[substrict_list[item]][:-1]
                                data['yongli_biaoshi'] = (basic_biaoshi + f'_{item+1}').replace('XQ','YL')
                                #常规填入
                                data['zhangjie'] = zhangjiehao
                                data['csx_mingcheng'] = zhangjieming
                                data['biaoshi'] = basic_biaoshi
                                data['zongsu'] = basic_zongshu
                                #步骤填入,首先根据substrict_list获取有几个步骤
                                
                                #循环行数
                                for x in range(temp_list[item+1]-(temp_list[item]+1)): #循环一个用例步骤预期数
                                    zuhe_dict = {"buzhou":"","yuqi":"","xuhao":""}
                                    #把每个步骤和预期都放进zuhe_dict中
                                    zuhe_dict["buzhou"] = buzhou_list[temp_list[item]+x+1]
                                    zuhe_dict["yuqi"] = yuqi_list[temp_list[item]+x+1]
                                    zuhe_dict["xuhao"] = str(x+1)
                                    data["zuhe"].append(zuhe_dict)
                                    
                                if item == 0:
                                    data['is_first'] = '1'
                                #判断是否为SU标题
                                fenge = data['biaoshi'].split("_")
                                #获取当前测试项分割后的个数
                                if fenge[-2] != is_fire_su:
                                    is_fire_su = fenge[-2]
                                    data['is_begin'] = "1"
                                    data['csxbs'] = zhuan_dict[fenge[-2]]
                                if self.parent.lineEdit.text():
                                    data['renyuan'] = self.parent.lineEdit.text()
                                #加入data_list
                                data_list.append(data)
                                yongli_count += 1 #用例计数加一
                                
                                self.sin_out.emit('###获取用例序号：{}'.format(yongli_count))
                                self.sin_out.emit('###该用例标识为：{}'.format(data['yongli_biaoshi']))
                        except:
                            self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格处理失败$$$$$$$$$$$$')
                            pass
            
        #关闭大纲文档（因为以及提取完毕）
        try:
            dagangfile.Close()
            self.w.Quit()
            pythoncom.CoUninitialize()
        except:
            self.sin_out.emit('function fail')
            self.w.Quit()
            pythoncom.CoUninitialize()
            return
    
        #打开模板文件进行渲染，然后就是用docxtpl生成用例
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "测试说明模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path) #模板导入成功
            
        except:
            QMessageBox.warning(self.parent,"出错了","导入模板出错请检查模板文件是否存在或名字不正确")
            return
        
        #开始渲染模板文件-有2层循环
        try:
            context = {
                "tables":data_list,
            }
            tpl.render(context)
            tpl.save("生成的说明文档.docx")
            QMessageBox.warning(self.parent,"生成文档成功","请查看当前工具根目录（生成的说明文档.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent,"生成文档出错","生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
        
##################################################################################
#大纲生成追踪关系
##################################################################################
class create_dagang_zhuisu(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        self.sin_out.emit("进入大纲追踪线程......")
        self.sin_out.emit("开始填写追踪......")
        
        #如果没有选择路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        #告诉windows单线程
        pythoncom.CoInitialize()
        #在用户选择的目录中查找大纲文档
        self.sin_out.emit('打开测评大纲文档...')
        
        #使用win32com打开-记得关闭
        #打开word应用
        self.w = DispatchEx('Word.Application')
        #self.w.visible=True
        self.w.DisplayAlerts = 0
        try:
            dagangfile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        curpath = Path.cwd() / 'need'
        zhuisu_path_tmp = curpath / 'document_templates' / '大纲追踪模板.docx'
        print(zhuisu_path_tmp)
        
        if zhuisu_path_tmp.is_file():
            self.sin_out.emit('已检测到有追溯模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return
        
        #创建个列表放数据
        data_list = []
        data2_list = []
        
        try:
            csx_tb_count = dagangfile.Tables.Count
            self.sin_out.emit('total:'+ str(csx_tb_count))
        except:
            self.sin_out.emit('不存在表格！')
            QMessageBox.warning(self.parent,'出错了','测试说明文档格式错误或者没有正确表格')
            try:
                dagangfile.Close()
            except:
                pass
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        for i in range(csx_tb_count):
            self.sin_out.emit(str(i))
            self.sin_out.emit("正在处理第{}个表格...".format(str(i)))
            print("正在处理第{}个表格...".format(str(i)))
            #准备填入的data
            data = {'xuqiu':[],'dg_zhangjie':'',
                    'mingcheng':'','biaoshi':''}
            data2 = {'xuqiu':[],'dg_zhangjie':'',
                    'mingcheng':'','biaoshi':''}
            if dagangfile.Tables[i].Rows.Count > 2:
                #注意win32com的Cell从1开始不是从0开始
                if dagangfile.Tables[i].Cell(1, 1).Range.Text.find('测试项名称') != -1:               
                    #一个用例不变内容获取
                    dagangfile.Tables[i].Rows.First.Select() #获取测试项章节号
                    zhangjiehao = self.w.Selection.Bookmarks("\headinglevel").\
                        Range.Paragraphs(1).Range.ListFormat.ListString #获取测试项章节名
                    zhangjieming = self.w.Selection.Bookmarks("\headinglevel").\
                        Range.Paragraphs(1).Range.Text.rstrip('\r')
                    biaoshi = dagangfile.Tables[i].Cell(1,4).Range.Text.rstrip()[:-2]
                    
                    #获取需规的章节号和描述
                    if dagangfile.Tables[i].Cell(2, 1).Range.Text.find("追踪关系") != -1:
                        zhuizong_tmp = dagangfile.Tables[i].Cell(2, 2).Range.Text[:-2]
                        #由于有/的存在，先判断/和隐含需求
                        zhuizong_list = zhuizong_tmp.split("\r")
                        print(zhuizong_list)
                        if zhuizong_tmp == "/" or zhuizong_tmp == "隐含需求":
                            xuqiu_dict = {'xq_zhangjie': '/', 'xq_miaoshu': '/'}
                            data['xuqiu'].append(xuqiu_dict)
                            data2['xuqiu'].append(xuqiu_dict)
                        else:
                            if len(zhuizong_list) >= 1:
                                for item in zhuizong_list:
                                    xuqiu_dict = {}
                                    if item.find("需求") != -1:
                                        try:
                                            match_string = re.search("\d(.\d)+", item).group()
                                            match_ming = item.split(match_string)[-1]
                                            xuqiu_dict['xq_zhangjie'] = match_string
                                            xuqiu_dict['xq_miaoshu'] = match_ming.lstrip(" ")
                                            data['xuqiu'].append(xuqiu_dict)
                                        except:
                                            self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格无章节号$$$$$$$$$$$$')
                                            self.sin_out.emit("转换终止！请检查表格中追踪关系有无章节号")
                                            pass
                                    else:
                                        try:
                                            match_string = re.search("\d(.\d)+", item).group()
                                            match_ming = item.split(match_string)[-1]
                                            xuqiu_dict['xq_zhangjie'] = match_string
                                            xuqiu_dict['xq_miaoshu'] = match_ming.lstrip(" ")
                                            data2['xuqiu'].append(xuqiu_dict)
                                        except:
                                            self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格无章节号$$$$$$$$$$$$')
                                            self.sin_out.emit("转换终止！请检查表格中追踪关系有无章节号")
                                            pass

                            #如果追踪关系行数小于1行的情况
                            else:
                                xuqiu_dict = {'xq_zhangjie': '/', 'xq_miaoshu': '/'}
                                data['xuqiu'].append(xuqiu_dict)
                                data2['xuqiu'].append(xuqiu_dict)

                        try:
                            data['dg_zhangjie'] = zhangjiehao
                            data['mingcheng'] = zhangjieming
                            data['biaoshi'] = biaoshi
                            data_list.append(data)
                            data2['dg_zhangjie'] = zhangjiehao
                            data2['mingcheng'] = zhangjieming
                            data2['biaoshi'] = biaoshi
                            data2_list.append(data2)
                        except:
                            print("获取追踪出错啦！")
                            self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格追踪处理失败$$$$$$$$$$$$')
                            pass
                          
        
        #最后关闭文档
        try:
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
        except:
            QMessageBox.warning(self.parent,"关闭文档失败","关闭文档失败！")
            return
        
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "大纲追踪模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path) #模板导入成功
            
        except:
            QMessageBox.warning(self.parent,"出错了","导入模板出错请检查模板文件是否存在或名字不正确")
            return
        
        #开始渲染模板文件
        try:
            context = {
                "tables":data_list,
                "tables2":data2_list,
            }
            tpl.render(context)
            tpl.save("生成的大纲追踪文档.docx")
            QMessageBox.warning(self.parent,"生成文档成功","请查看当前工具根目录（生成的大纲追踪文档.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent,"生成文档出错","生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
        
##################################################################################
#单元测试UAS转换
##################################################################################
class create_danyuan(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        self.sin_out.emit("进入单元测试SunwiseAUnit转换线程......")
        self.sin_out.emit("开始填写文档......")
         
        #如果没有选择路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        
        #告诉windows单线程
        pythoncom.CoInitialize()
        #在用户选择的目录中查找UAS单位测试报告文档
        self.sin_out.emit('打开单元测试原文件...')
        
        #使用win32com打开-记得关闭
        #打开word应用
        self.w = DispatchEx('Word.Application')
        #self.w.visible=True
        self.w.DisplayAlerts = 0
        try:
            danyuanfile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        curpath = Path.cwd() / 'need'
        danyuan_file_tmp = curpath / 'document_templates' / 'SunwiseAUnit单元测试转换模板.docx'
        print(danyuan_file_tmp)
        
        if danyuan_file_tmp.is_file():
            self.sin_out.emit('已检测到有追溯模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return
        
        #创建个列表放数据-important
        data_list = []
        
        #try统计表格数量
        try:
            csx_tb_count = danyuanfile.Tables.Count
            self.sin_out.emit('total:'+ str(csx_tb_count))
            self.sin_out.emit("正在调用word文档操作接口,可能会有点慢...")
        except:
            self.sin_out.emit('不存在表格！')
            QMessageBox.warning(self.parent,'出错了','测试说明文档格式错误或者没有正确表格')
            try:
                danyuanfile.Close()
            except:
                pass
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        #开始处理表格-important
        #我先统计有多少个生成的表格-即用例有多少个呗
        yongli_count = 0
        for i in range(csx_tb_count):
            if danyuanfile.Tables[i].Rows.Count > 2:
                #注意win32com的Cell从1开始不是从0开始
                if danyuanfile.Tables[i].Cell(1, 1).Range.Text.find('用例名称') != -1: 
                    yongli_count += 1
                    
        #yongli_num指向当前处理的用例
        yongli_num = 0
        hanshuming = ''
        hanshuming_duibi = ''
        wenjian = ''
        wenjian_duibi = ''
        for i in range(csx_tb_count):
            self.sin_out.emit('正在处理的表格序号：' + str(yongli_num + 1))
            self.sin_out.emit(str(i))
            #准备填入的data
            data = {'ruanjian_ming':'','ruanjian_biaoshi':'yongli_biaoshi','wenjian_ming':'',\
                'hanshu_ming':'','bianlian_and_canshu':'','zhuang':[],\
                    'yuqi_jieguo':'','ceshi_jieguo':'','is_begin':'0','is_wenjian':'0'}
            
            #填入用户输入的软件名
            try:
                data['ruanjian_ming'] = self.parent.lineEdit_2.text()
                data['ruanjian_biaoshi'] = self.parent.lineEdit_3.text()
                
            except:
                QMessageBox.critical(self.parent,"未填入数据","请先填入软件名和软件标识或.C名称")
                self.w.Quit()
                pythoncom.CoUninitialize()
                self.parent.tabWidget.setEnabled(True)
                pass
            
            #找到函数名
            if danyuanfile.Tables[i].Rows.Count > 2:
                if danyuanfile.Tables[i].Cell(1, 1).Range.Text.find('功能描述') != -1: 
                    danyuanfile.Tables[i].Cell(1, 1).Range.Select()
                    self.w.Selection.MoveUp()
                    self.w.Selection.MoveUp()
                    self.w.Selection.MoveUp()
                    s = self.w.Selection.Paragraphs(1).Range.Text[:-1]
                    s1 = s.split(". ")[-1]
                    #放入函数名比对
                    if s1 != hanshuming_duibi:
                        hanshuming_duibi = s1
                        
                    #再向上看2行
                    self.w.Selection.MoveUp()
                    self.w.Selection.MoveUp()
                    temp = self.w.Selection.Paragraphs(1).Range.Text[:2]
                    temp2 = self.w.Selection.Paragraphs(1).Range.Text[:-1]
                    s2 = temp2.split(". ")[-1]
                    if temp == "文件":
                        if s2 != wenjian_duibi:
                            wenjian_duibi = s2

            #找章节号
                    
            
            if danyuanfile.Tables[i].Rows.Count > 2:
                #注意win32com的Cell从1开始不是从0开始
                if danyuanfile.Tables[i].Cell(1, 1).Range.Text.find('用例名称') != -1:  
                    #TODO：如何找到测试模块？
                    biaoshi_temp = danyuanfile.Tables[i].Cell(1, 4).Range.Text[:-2]
                    data['yongli_biaoshi'] = biaoshi_temp
                    
                    #获取表格中参数组合()
                    quanju = danyuanfile.Tables[i].Cell(5, 3).Range.Text[:-2]
                    hcan = danyuanfile.Tables[i].Cell(6, 3).Range.Text[:-2]
                    qitashu = danyuanfile.Tables[i].Cell(7, 3).Range.Text[:-2]
                    
                    if quanju.find('无') != -1: 
                        quanju = ""
                    if hcan.find('无') != -1: 
                        hcan = ""
                    if qitashu.find('无') != -1: 
                        qitashu = ""
                    
                    data['bianlian_and_canshu'] = quanju + hcan + qitashu
                    #将预期结果和测试结果填入
                    data['yuqi_jieguo'] = danyuanfile.Tables[i].Cell(8, 2).Range.Text[:-2]
                    data['ceshi_jieguo'] = danyuanfile.Tables[i].Cell(13, 2).Range.Text[:-2]
                    #函数名获取
                    if hanshuming_duibi != hanshuming:
                        hanshuming = hanshuming_duibi
                        data['is_begin'] = '1'
                    data['hanshu_ming'] = hanshuming_duibi
                    #文件名获取
                    if wenjian_duibi != wenjian:
                        wenjian = wenjian_duibi
                        data['is_wenjian'] = '1'
                    data['wenjian_ming'] = wenjian_duibi
                    
                    data_list.append(data)
                    yongli_num += 1 #用例创建加一
                    
                elif danyuanfile.Tables[i].Cell(1, 2).Range.Text.find('定义') != -1: 
                    #定义个桩函数dict
                        zhuang_dict = {'zhuang_name':'','zhuang_dingyi':'','zhuang_fanhui':'',\
                            'zhuang_fuzuoyong':''}
                        
                        zhuang_dict['zhuang_name'] = danyuanfile.Tables[i].Cell(1, 1).Range.Text[:-2]
                        zhuang_dict['zhuang_dingyi'] = danyuanfile.Tables[i].Cell(1, 3).Range.Text[:-2]
                        zhuang_dict['zhuang_fanhui'] = danyuanfile.Tables[i].Cell(2, 3).Range.Paragraphs(1).\
                            Range.Text[:-1]
                        
                        #副作用可能有多行
                        fuzuoyong_temp = ''
                        for count_fuzuo in range(len(danyuanfile.Tables[i].Cell(2, 3).Range.Paragraphs)-2):
                            fuzuoyong_temp = fuzuoyong_temp + ';' + danyuanfile.Tables[i].Cell(2, 3).\
                            Range.Paragraphs(count_fuzuo + 3).Range.Text[:-2].replace(" ", "")
                        zhuang_dict['zhuang_fuzuoyong'] = fuzuoyong_temp
                        
                        data_list[yongli_num - 1]['zhuang'].append(zhuang_dict)
                    
                    #气死了这里要写成[:-2]而不是[-2]!
                          
        #最后关闭文档
        try:
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
        except:
            QMessageBox.warning(self.parent,"关闭文档失败","关闭文档失败！")
            return
        
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "SunwiseAUnit单元测试转换模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path) #模板导入成功
        except:
            QMessageBox.warning(self.parent,"出错了","导入模板出错请检查模板文件是否存在或名字不正确")
            return
        
        #开始渲染模板文件
        try:
            context = {
                "tables":data_list,
            }
            tpl.render(context)
            tpl.save("软件单元测试用例记录表.docx")
            QMessageBox.warning(self.parent,"生成文档成功","请查看当前工具根目录（软件单元测试用例记录表.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent,"生成文档出错","生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
 
##################################################################################
#测试说明追踪以及用例表
################################################################################## 
class create_shuoming_zhuisu(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self, parent):
        super().__init__()
        self.parent = parent

    def run(self):
        self.sin_out.emit("进入说明追踪线程......")
        self.sin_out.emit("开始填写说明追踪以及用例表格......")

        # 如果没有选择文件
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        # 告诉windows单线程
        pythoncom.CoInitialize()
        # 在用户选择的目录中查找大纲文档
        self.sin_out.emit('打开测试说明文档...')

        # 使用win32com打开-记得关闭
        # 打开word应用
        self.w = DispatchEx('Word.Application')
        # self.w.visible=True
        self.w.DisplayAlerts = 0
        try:
            shuomingfile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        self.sin_out.emit('已正确打开说明文档...')
        curpath = Path.cwd() / 'need'
        zhuisu_path_tmp = curpath / 'document_templates' / '说明追踪模板.docx'
        print("打开追踪模板文件",zhuisu_path_tmp)

        if zhuisu_path_tmp.is_file():
            self.sin_out.emit('已检测到有说明追溯模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return

        # 创建个列表放数据
        data_list = []
        data2_list = []

        #统计整个表格数量用于processbar显示进度
        try:
            tb_count = shuomingfile.Tables.Count
            self.sin_out.emit('total:' + str(tb_count))
        except:
            self.sin_out.emit('不存在表格！')
            QMessageBox.warning(self.parent, '出错了', '测试说明文档格式错误或者没有正确表格')
            try:
                shuomingfile.Close()
            except:
                QMessageBox.warning(self.parent, '错误', "未正确关闭Word文档！")
                return
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return

        #遍历循环表格，这里面就要初始化数据dict了
        #不能像大纲追踪一样data在循环表格里面
        #创建一个大纲测试项索引
        csx_name = ''
        data = {'dg_zhangjie': '', 'mingcheng': '','biaoshi': '', 'yongli':[],'index':0}
        for i in range(tb_count):
            self.sin_out.emit(str(i))
            self.sin_out.emit("正在处理第{}个表格...".format(str(i+1)))
            print("正在处理第{}个表格...".format(str(i+1)))
            # 准备填入的data
            data2 = {'yongli_ming':'','yongli_biaoshi':'','yongli_zongsu':''}
            yongli_dict = {'yongli_ming':'','yongli_biaoshi':''}
            yongliming = ''
            biaoshi = ''
            zongsu = ''
            zhui_temp = ''
            
            if shuomingfile.Tables[i].Rows.Count > 2:
                # 注意win32com的Cell从1开始不是从0开始
                if shuomingfile.Tables[i].Cell(1, 1).Range.Text.find('测试用例名称') != -1 or \
                    shuomingfile.Tables[i].Cell(2, 1).Range.Text.find('测试用例名称') != -1:
                    try:    
                        #取出cell(1,,1)的数据
                        table_heard = shuomingfile.Tables[i].Cell(1, 1).Range.Text
                        if table_heard.find("测试用例名称") != -1:
                            yongliming = shuomingfile.Tables[i].Cell(1, 2).Range.Text.rstrip()[:-2]
                            biaoshi = shuomingfile.Tables[i].Cell(1, 4).Range.Text.rstrip()[:-2]
                            zongsu = shuomingfile.Tables[i].Cell(3, 2).Range.Text.rstrip()[:-2]
                            zhui_temp = shuomingfile.Tables[i].Cell(2,2).Range.Text.rstrip()[:-2]
                        elif table_heard.find('用例') != -1:
                            yongliming = shuomingfile.Tables[i].Cell(2, 2).Range.Text.rstrip()[:-2]
                            biaoshi = shuomingfile.Tables[i].Cell(2, 4).Range.Text.rstrip()[:-2]
                            zongsu = shuomingfile.Tables[i].Cell(4, 2).Range.Text.rstrip()[:-2]
                            zhui_temp = shuomingfile.Tables[i].Cell(3,2).Range.Text.rstrip()[:-2]
                        else:
                            self.sin_out.emit("未找到合适的填写数据，退出处理")
                            print("未找到合适的填写数据")
                            self.w.Quit()
                            pythoncom.CoUninitialize()
                            self.parent.tabWidget.setEnabled(True)
                            return
                        shuomingfile.Tables[i].Rows.First.Select()  # 获取测试项章节号
                        #############################目前模板不用获取用例章节号暂时省去
                        # zhangjiehao = self.w.Selection.Bookmarks("\headinglevel"). \
                        #     Range.Paragraphs(1).Range.ListFormat.ListString  # 获取测试项章节名
                        ##############################################################
                        zhangjieming = self.w.Selection.Bookmarks("\headinglevel"). \
                            Range.Paragraphs(1).Range.Text.rstrip('\r')
                        
                        yongli_dict['yongli_ming'] = yongliming
                        yongli_dict['yongli_biaoshi'] = biaoshi
                        data2['yongli_ming'] = yongliming
                        data2['yongli_biaoshi'] = biaoshi
                        data2['yongli_zongsu'] = zongsu
                        data2_list.append(data2)
                        print("当前yongli_dict为：",yongli_dict)
                        # 获取大纲的章节号和用例名，而且data按自己的来
                        ## 按python行进行分割为列表
                        zhui_list = zhui_temp.split("\r")
                        if len(zhui_list) == 3:
                            if zhui_list[1].find("需求") != -1:
                                #使用re模块正则表达式
                                match_string = re.search("\d(.\d)+",zhui_list[1]).group()
                                match_ming = zhui_list[1].split(match_string)[-1]
                                #使用re.sub模块替换为空
                                rules = "[)(）（] "
                                match_ming = re.sub(rules,'',match_ming)
                                if zhui_list[2]:
                                    rules = ":"
                                    dg_biaoshi_temp = re.sub(rules,'：',zhui_list[2])
                                    dg_biaoshi = dg_biaoshi_temp.split("：")[-1]
                                #判断是否是新的测试项，如果是新的索引index加1，创建新dict进入
                                if zhangjieming == csx_name:
                                    data['yongli'].append(yongli_dict)
                                #如果测试项是新的
                                else:
                                    data['yongli'].append(yongli_dict)
                                    data['dg_zhangjie'] = match_string
                                    data['mingcheng'] = match_ming
                                    data['biaoshi'] = dg_biaoshi
                                    data_list.append(data)
                                    data_index = data['index'] + 1
                                    csx_name = zhangjieming
                                    #清空data数据
                                    data = {'dg_zhangjie': '', 'mingcheng': '','biaoshi': '', 'yongli':[],'index':data_index}
                                    self.sin_out.emit("已处理第{}个测试项...".format(data['index']))
                                        
                    except:
                        self.sin_out.emit(f'$$$$$$$$$$$$第{str(i+1)}个表格，获取单元格内容不存在$$$$$$$$$$$$')
                        pass
        # 最后关闭文档
        try:
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
        except:
            QMessageBox.warning(self.parent, "关闭文档失败", "关闭文档失败！")
            return

        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "说明追踪模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path)  # 模板导入成功

        except:
            QMessageBox.warning(self.parent, "出错了", "导入模板出错请检查模板文件是否存在或名字不正确")
            return

        # 开始渲染模板文件
        try:
            context = {
                "tables": data_list,
                "tables2": data2_list,
            }
            tpl.render(context)
            tpl.save("说明追踪文档.docx")
            QMessageBox.about(self.parent, "生成文档成功", "请查看当前工具根目录（说明追踪文档.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent, "生成文档出错", "生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
        
##################################################################################
#根据说明生成测试记录线程
##################################################################################
class create_jilu(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        #用来储存章节号转换
        zhuan_dict = {'DC':'文档审查','SU':'功能测试','CR':'代码审查','SA':'静态分析','AC':'性能测试',\
            'IO':'接口测试','SE':'安全性测试','BT':'边界测试','RE':'恢复性测试','ST':'强度测试',\
                'AT':'余量测试','GUI':'人机交互界面测试','DP':'数据处理测试','JR':'兼容性测试',\
                    'LG':'逻辑测试'}
        
        self.sin_out.emit("进入根据说明转换记录线程......")
        self.sin_out.emit("开始转换......")
        #如果没有选择文件则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        pythoncom.CoInitialize()
        self.sin_out.emit('打开说明文档...')
        self.w = DispatchEx('Word.Application')
        self.w.DisplayAlerts = 0
        try:
            shuomingfile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        self.sin_out.emit('复制测试记录文档模板到本程序所在目录...')
        curpath = Path.cwd() / 'need'
        shuoming_path_tmp = curpath / 'document_templates' / '说明生成记录模板.docx'
        if shuoming_path_tmp.is_file():
            self.sin_out.emit('已检测到有记录模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return
        #创建一个字典来储存单个用例
        data_list = []
        #获取表格数量
        try:
            csx_tb_count = shuomingfile.Tables.Count
            self.sin_out.emit('total:'+ str(csx_tb_count))
        except:
            self.sin_out.emit('不存在表格，请检查文档！')
            QMessageBox.warning(self.parent,'出错了','测试说明文档格式错误或者没有正确表格')
            try:
                shuomingfile.Close()
            except:
                pass
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        #用来储存章节号中的DC、SU等标识，用于章节号判断
        is_type_su = ""
        #储存章节号标志
        is_fire_su = ""
        for i in range(csx_tb_count):
            self.sin_out.emit(str(i))
            self.sin_out.emit("正在处理第{}个表格".format(i+1))
            if shuomingfile.Tables[i].Rows.Count > 2:
                if shuomingfile.Tables[i].Cell(2, 1).Range.Text.find('测试用例名称') != -1:               
                    #一个用例不变内容获取
                    try:
                        shuomingfile.Tables[i].Rows.First.Select()
                        #获取章节名，用于判断
                        zhangjieming = self.w.Selection.Bookmarks("\headinglevel").\
                            Range.Paragraphs(1).Range.Text.rstrip('\r')
                        #获取表格基本信息
                        mingcheng = shuomingfile.Tables[i].Cell(2,2).Range.Text[:-2]
                        biaoshi = shuomingfile.Tables[i].Cell(2,4).Range.Text[:-2]
                        self.sin_out.emit(f"正在处理{biaoshi}用例{mingcheng}")
                        zhuizong = shuomingfile.Tables[i].Cell(3,2).Range.Text[:-2]
                        zongsu = shuomingfile.Tables[i].Cell(4,2).Range.Text[:-2]
                        chushihua = shuomingfile.Tables[i].Cell(5,2).Range.Text[:-2]
                        qianti = shuomingfile.Tables[i].Cell(6,2).Range.Text[:-2]

                        #缓存一个data数据
                        data = {'mingcheng':'','biaoshi':'','zhuizong':'','is_first':'0',\
                            'zongsu':'','chushihua':'','qianti':'','zuhe':[],'is_begin':'0',\
                                'csx_type':'','csx_name':''}
                        #获取步骤和预期
                        step_count = shuomingfile.Tables[i].Rows.Count - 11
                        for j in range(step_count):
                            buzhou_dict = {'buzhou':"",'yuqi':"",'xuhao':''}
                            buzhou_dict['buzhou'] = shuomingfile.Tables[i].Cell(j+9,2).Range.Text[:-2]
                            buzhou_dict['yuqi'] = shuomingfile.Tables[i].Cell(j+9,3).Range.Text[:-2]
                            buzhou_dict['xuhao'] = str(j+1)
                            data['zuhe'].append(buzhou_dict)
                            
                        #开始判断当前是否为测试项的第一个，如果是第一个则is_first改为1
                        if is_fire_su != zhangjieming:
                            is_fire_su = zhangjieming
                            data['is_first'] = '1'
                            
                        #判断测试类型，这里从标识里面获取
                        biaoshi_list = biaoshi.split("_")
                        print('当前取的类型列表分割：',biaoshi_list)
                        if len(biaoshi_list) >= 4:
                            biaoshi_tmp = biaoshi_list[-3]
                        else:
                            biaoshi_tmp = biaoshi_list[1]
                        if biaoshi_tmp != is_type_su:
                            is_type_su = biaoshi_tmp
                            data['is_begin'] = '1'
                            if zhuan_dict[biaoshi_tmp] == '文档审查' or  zhuan_dict[biaoshi_tmp] == '代码审查' or \
                                zhuan_dict[biaoshi_tmp] == '静态分析':
                                data['is_first'] = '0'
                            
                        #data补全
                        data['mingcheng'] = mingcheng
                        data['biaoshi'] = biaoshi
                        data['zhuizong'] = zhuizong.replace('\r','\n')
                        data['zongsu'] = zongsu
                        data['chushihua'] = chushihua
                        data['qianti'] = qianti
                        data['csx_type'] = zhuan_dict[biaoshi_tmp]
                        data['csx_name'] = zhangjieming
                        data_list.append(data)
                        self.sin_out.emit("处理完毕{}用例".format(biaoshi))
                    except:
                        self.sin_out.emit(f'错误！第{i+1}个表格处理失败!')
                    
                    
        #关闭大纲文档（因为以及提取完毕）
        try:
            shuomingfile.Close()
            self.w.Quit()
            pythoncom.CoUninitialize()
        except:
            self.sin_out.emit('function fail')
            self.w.Quit()
            pythoncom.CoUninitialize()
            return
    
        #打开模板文件进行渲染，然后就是用docxtpl生成用例
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "说明生成记录模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path) #模板导入成功
            
        except:
            QMessageBox.warning(self.parent,"出错了","导入模板出错请检查模板文件是否存在或名字不正确")
            return
        
        #开始渲染模板文件-有2层循环
        try:
            context = {
                "tables":data_list,
                "cs_renyuan":self.parent.lineEdit_4.text(),
                "jc_renyuan":self.parent.lineEdit_5.text(),
                "shijian":self.parent.lineEdit_6.text(),
            }
            tpl.render(context)
            tpl.save("生成的测试记录文档.docx")
            QMessageBox.warning(self.parent,"生成文档成功","请查看当前工具根目录（生成的测试记录文档.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent,"生成文档出错","生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
    
##################################################################################
#根据测试记录反向生成说明
##################################################################################
class create_shuomingfanxiang(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        #用来储存测试项DC等转换
        zhuan_dict = {'DC':'文档审查','SU':'功能测试','CR':'代码审查','SA':'静态分析','AC':'性能测试',\
            'IO':'接口测试','SE':'安全性测试','BT':'边界测试','RE':'恢复性测试','ST':'强度测试',\
                'AT':'余量测试','GUI':'人机交互界面测试','DP':'数据处理测试','JR':'兼容性测试',\
                    'LG':'逻辑测试'}
        
        self.sin_out.emit("进入测试记录转说明......")
        self.sin_out.emit("开始转换......")
        #如果没有选择文件路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        pythoncom.CoInitialize()
        self.sin_out.emit('打开测试记录文件...')
        self.w = DispatchEx('Word.Application')
        #self.w.visible=True
        self.w.DisplayAlerts = 0
        try:
            jilufile = self.w.Documents.Open(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        self.sin_out.emit('复制测试说明文档模板到本程序所在目录...')
        curpath = Path.cwd() / 'need'
        shuoming_path_tmp = curpath / 'document_templates' / '反向测试说明模板.docx'
        print(shuoming_path_tmp)
        if shuoming_path_tmp.is_file():
            self.sin_out.emit('已检测到有说明模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return
        #创建一个字典来储存单个用例
        data_list = []
        #获取表格数量
        try:
            csx_tb_count = jilufile.Tables.Count
            self.sin_out.emit('total:'+ str(csx_tb_count))
        except:
            self.sin_out.emit('不存在表格！')
            QMessageBox.warning(self.parent,'出错了','测试说明文档格式错误或者没有正确表格')
            try:
                jilufile.Close()
            except:
                pass
            self.w.Quit()
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        #初始化表格外全局变量
        is_fire_su = ''
        is_type_su = ''
        for i in range(csx_tb_count):
            self.sin_out.emit(str(i))
            self.sin_out.emit(f"正在处理第{str(i+1)}个表格")
            if jilufile.Tables[i].Rows.Count > 2:
                if jilufile.Tables[i].Cell(2, 1).Range.Text.find('测试用例名称') != -1:
                    #将表格中信息全部先拿出来
                    try:
                        jilufile.Tables[i].Rows.First.Select()
                        zhangjieming = self.w.Selection.Bookmarks("\headinglevel").\
                            Range.Paragraphs(1).Range.Text.rstrip('\r')
                        zhangjiehao = self.w.Selection.Bookmarks("\headinglevel").\
                            Range.Paragraphs(1).Range.ListFormat.ListString
                        mingcheng = jilufile.Tables[i].Cell(2, 2).Range.Text[:-2]
                        biaoshi = jilufile.Tables[i].Cell(2, 4).Range.Text[:-2]
                        self.sin_out.emit(f"正在处理{biaoshi}-用例{mingcheng}")
                        zhuizong = jilufile.Tables[i].Cell(3,2).Range.Text[:-2]
                        zongsu = jilufile.Tables[i].Cell(4,2).Range.Text[:-2]
                        chushi = jilufile.Tables[i].Cell(5,2).Range.Text[:-2]
                        qianti = jilufile.Tables[i].Cell(6,2).Range.Text[:-2]
                        
                        #缓存一个data数据
                        data = {'mingcheng':'','biaoshi':'','zhuizong':'','is_first':'0',\
                            'zongsu':'','chushi':'','qianti':'','zuhe':[],'is_begin':'0',\
                                'csx_type':'','csx_mingcheng':'','renyuan':''}
                        #获取步骤和预期
                        step_count = jilufile.Tables[i].Rows.Count - 12
                        #获取人员信息
                        data['renyuan'] = jilufile.Tables[i].Cell(10+step_count,2).Range.Text[:-2]
                        for j in range(step_count):
                            buzhou_dict = {'buzhou':"",'yuqi':"",'xuhao':''}
                            buzhou_dict['buzhou'] = jilufile.Tables[i].Cell(j+9,2).Range.Text[:-2]
                            buzhou_dict['yuqi'] = jilufile.Tables[i].Cell(j+9,3).Range.Text[:-2]
                            buzhou_dict['xuhao'] = str(j+1)
                            data['zuhe'].append(buzhou_dict)
                        
                        #开始判断当前是否为测试项的第一个，如果是第一个则is_first改为1
                        if is_fire_su != zhangjieming:
                            is_fire_su = zhangjieming
                            data['is_first'] = '1'
                        #判断测试类型，这里从标识里面获取
                        biaoshi_list = biaoshi.split("_")
                        print('当前取的类型列表分割：',biaoshi_list)
                        if len(biaoshi_list) >= 4:
                            biaoshi_tmp = biaoshi_list[-3]
                        else:
                            biaoshi_tmp = biaoshi_list[1]
                        if biaoshi_tmp != is_type_su:
                            is_type_su = biaoshi_tmp
                            data['is_begin'] = '1'
                            if zhuan_dict[biaoshi_tmp] == '文档审查' or  zhuan_dict[biaoshi_tmp] == '代码审查' or \
                                zhuan_dict[biaoshi_tmp] == '静态分析':
                                data['is_first'] = '0'
                        
                        #data补全
                        data['mingcheng'] = mingcheng
                        data['biaoshi'] = biaoshi
                        data['zhuizong'] = zhuizong.replace('\r','\n')
                        data['zongsu'] = zongsu
                        data['chushi'] = chushi
                        data['qianti'] = qianti
                        data['csx_type'] = zhuan_dict[biaoshi_tmp]
                        data['csx_mingcheng'] = zhangjieming
                        data_list.append(data)
                        self.sin_out.emit("处理完毕{}用例".format(biaoshi))
                    except:
                        self.sin_out.emit("第{}个表格处理失败，请检查".format(str(i+1)))
                        pass
                        
                else:
                    self.sin_out.emit("该表格生成错误，请检查是否存在用例序号，每个用例必须有序号且必须包含【记录】两个字...")
            else:
                self.sin_out.emit("该表格生成错误，请检查表格是否存在并大于2行...")  
            
        #关闭大纲文档（因为以及提取完毕）
        try:
            jilufile.Close()
            self.w.Quit()
            pythoncom.CoUninitialize()
        except:
            self.sin_out.emit('function fail')
            self.w.Quit()
            pythoncom.CoUninitialize()
            return
    
        #打开模板文件进行渲染，然后就是用docxtpl生成用例
        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / "反向测试说明模板.docx"
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path) #模板导入成功
            
        except:
            QMessageBox.warning(self.parent,"出错了","导入模板出错请检查模板文件是否存在或名字不正确")
            return
        
        #开始渲染模板文件-有2层循环
        try:
            context = {
                "tables":data_list,
            }
            tpl.render(context)
            tpl.save("反向生成的说明文档.docx")
            QMessageBox.warning(self.parent,"生成文档成功","请查看当前工具根目录（反向生成的说明文档.docx）,【注意】生成\
                    的文件章节号中存在错误，请自行添加二级章节号，并且将三级章节号降级处理")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent,"生成文档出错","生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return
    
##################################################################################
#自动填充单元格线程
##################################################################################
class create_zidong(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent

    def run(self):
        self.sin_out.emit('开始...')
        if self.parent.open_file_name == '':
            self.sin_out.emit('请点击“选择文档”按钮选择要填充的文档')
            self.parent.tabWidget.setEnabled(True)
            QMessageBox.warning(self.parent,'出错了！','请选择要填充的文档！')
            return
        try:
            t_s_file = docx.Document(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.parent.tabWidget.setEnabled(True)
            QMessageBox.warning(self.parent,'出错了！','打开选择的文档失败，请确认文档类型为docx，且未被打开！')
            return
        if self.parent.lineEdit_9.text() == '':
            self.sin_out.emit('单元格左侧不能为空！!!!')
            self.parent.tabWidget.setEnabled(True)
            QMessageBox.warning(self.parent,'出错了！','单元格标题不能为空！')
            return
        if self.parent.lineEdit_10.text() == '':
            self.sin_out.emit('确定填充内容为空吗？填充内容为空相当于清空操作。可直接点击清空按钮！！')
            self.parent.tabWidget.setEnabled(True)
            QMessageBox.warning(self.parent,'警告！','确定填充内容为空吗？填充内容为空相当于清空操作。可点击清空按钮！！')
            return
        tmp_fill = self.parent.lineEdit_11.text()
        if tmp_fill == '':
            tmp_fill = str(len(t_s_file.tables))

        if (tmp_fill.strip().isdigit()) and (int(tmp_fill.strip()) < len(
                t_s_file.tables)):
            tmp_ran = int(tmp_fill)
        else:
            tmp_ran = len(t_s_file.tables)

        tmp_fillnum = 0
        k = 0

        self.sin_out.emit('total:' + str(tmp_ran))

        self.parent.progressBar.setRange(0,tmp_ran-1)
        for ft1 in t_s_file.tables:
            k += 1
            self.sin_out.emit(str(k))
            self.parent.progressBar.setValue(k)
            tmp_row = 0
            for r in ft1.rows:
                tmpflag = 0
                tmp_column = 0
                for cell in r.cells:
                    if cell.text.strip() == self.parent.lineEdit_9.text():
                        while ft1.cell(tmp_row, tmp_column).text.strip() == self.parent.lineEdit_9.text():
                            tmp_column += 1
                        #这里如果需要替换还是不替换
                        if ft1.cell(tmp_row, tmp_column).text == '':
                            ft1.cell(tmp_row, tmp_column).text = self.parent.lineEdit_10.text()
                            tmp_fillnum += 1
                        else:
                            pass
                        tmpflag = 1
                        break
                    else:
                        tmp_column += 1
                if tmpflag == 1:
                    break
                tmp_row += 1
            if tmp_fillnum >= int(tmp_fill):
                break

        try:
            t_s_file.save(self.parent.open_file_name[0])
            self.sin_out.emit('function success')
            self.parent.tabWidget.setEnabled(True)
            return
        except:
            self.parent.tabWidget.setEnabled(True)
            self.sin_out.emit('function fail')
            QMessageBox.information(self.parent,'','填充完成！')
            return

##################################################################################
#清空单元格线程
##################################################################################
class clear_cell(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent

    def run(self):
        self.sin_out.emit('开始...')
        if self.parent.open_file_name == '':
            self.sin_out.emit('请点击“选择文档”按钮选择要填充的文档')
            self.parent.tabWidget.setEnabled(True)
            QMessageBox.warning(self.parent,'出错了！','请选择要填充的文档！')
            return
        try:
            t_s_file = docx.Document(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:要填充的文档')
            QMessageBox.warning(self.parent,'出错了！','打开选择的文档失败，请确认文档类型为docx，且未被打开！')
            self.parent.tabWidget.setEnabled(True)
            return
        if self.parent.lineEdit_9.text() == '':
            self.parent.tabWidget.setEnabled(True)
            self.sin_out.emit('单元格标题不能为空！')
            QMessageBox.warning(self.parent,'出错了！','单元格标题不能为空！')
            return
        tmp_tblcnt = len(t_s_file.tables)
        k = 0
        self.sin_out.emit('total:' + str(tmp_tblcnt))
        for ft1 in t_s_file.tables:
            k += 1
            self.sin_out.emit(str(k))
            tmp_row = 0
            for r in ft1.rows:
                tmpflag = 0
                tmp_column = 0
                for cell in r.cells:

                    if cell.text.strip() == self.parent.lineEdit_9.text():
                        while ft1.cell(tmp_row, tmp_column).text.strip(
                        ) == self.parent.lineEdit_9.text():
                            tmp_column += 1

                        ft1.cell(tmp_row, tmp_column).text = ''
                        tmpflag = 1
                        break
                    else:
                        tmp_column += 1
                if tmpflag == 1:
                    break
                tmp_row += 1
        try:
            t_s_file.save(self.parent.open_file_name[0])
            self.sin_out.emit('function success')
            self.parent.tabWidget.setEnabled(True)
            return
        except:
            self.parent.tabWidget.setEnabled(True)
            self.sin_out.emit('function fail')
            QMessageBox.information(self.parent,'','清空单元格成功！')
            return
        
##################################################################################
#提取表格内容线程
##################################################################################
class get_content(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent

    #获取文档中表格内容函数
    def run(self):
        curpath = Path.cwd()
        content_tmp = curpath / 'need' / 'document_templates' / 'get_content.docx'
        shutil.copy(content_tmp, curpath)
        content_tmp_path = curpath / 'get_content.docx'
        print(content_tmp_path)
        try:
            #c_file = self.w.Documents.Add()
            c_file = docx.Document(content_tmp_path)
        except:
            self.sin_out.emit('open failed:文档模板')
            self.parent.tabWidget.setEnabled(True)
            return

        try:
            s_file = docx.Document(self.parent.open_file_name[0])
            s_tbls = s_file.tables
        except:
            self.sin_out.emit('open failed:选择的文档')
            c_file.save(content_tmp_path)
            self.parent.tabWidget.setEnabled(True)
            return
        #原来是5,6,7
        if self.parent.lineEdit_12.text() == '' and self.parent.lineEdit_13.text() == '' and self.parent.lineEdit_14.text() == '':

            self.sin_out.emit('warning:请至少填写一个要提取的内容的标题，\n标题为要提取的单元格的前一单元格中的内容!')
            c_file.save(content_tmp_path)
            s_file.save(self.parent.open_file_name[0])
            self.parent.tabWidget.setEnabled(True)
            return

        line_list = [
            self.parent.lineEdit_12.text(),
            self.parent.lineEdit_13.text(),
            self.parent.lineEdit_14.text()
        ]
        self.sin_out.emit('开始提取...')
        rownum = 0
        self.sin_out.emit('total:' + str(len(s_tbls)))
        for stb in s_tbls:
            c_file.tables[0].add_row()
            rownum += 1
            self.sin_out.emit(str(rownum))

            row = 0
            for r1 in stb.rows:
                col = 0
                for ce in r1.cells:
                    if line_list[0] != '' and ce.text == line_list[0]:
                        while stb.cell(row, col).text == line_list[0]:
                            col += 1
                        c_file.tables[0].cell(rownum,
                                              0).text = stb.cell(row, col).text
                        break
                    col += 1
                col = 0
                for ce in r1.cells:
                    if line_list[1] != '' and ce.text == line_list[1]:
                        while stb.cell(row, col).text == line_list[1]:
                            col += 1
                        c_file.tables[0].cell(rownum,
                                              1).text = stb.cell(row, col).text
                        break
                    col += 1
                col = 0
                for ce in r1.cells:
                    if line_list[2] != '' and ce.text == line_list[2]:
                        while stb.cell(row, col).text == line_list[2]:
                            col += 1
                        c_file.tables[0].cell(rownum,
                                              2).text = stb.cell(row, col).text
                        break
                    col += 1
                row += 1
        try:
            c_file.save(content_tmp_path)
            s_file.save(self.parent.open_file_name[0])
            self.sin_out.emit('function success')
            self.sin_out.emit('生成文件名为(get_content.docx)，在根目录下查看')
            self.parent.tabWidget.setEnabled(True)
            return
        except:
            self.sin_out.emit('function fail')
            self.parent.tabWidget.setEnabled(True)
            return

##################################################################################
#测评报告追溯表
################################################################################## 
class create_baogao_zhuisu(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self, parent):
        super().__init__()
        self.parent = parent

    def run(self):
        self.sin_out.emit("进入报告追溯线程......")
        self.sin_out.emit("开始填写报告追溯表......")

        # 如果没有选择文件
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        self.sin_out.emit('打开测试记录文档...')
        # 打开word应用
        try:
            doc_path = self.parent.open_file_name[0]
            doc = Document(doc_path)
        except:
            self.sin_out.emit('open failed:选择的文档')
            self.parent.tabWidget.setEnabled(True)
            return
        
        self.sin_out.emit('已正确打开说明文档...')
        curpath = Path.cwd() / 'need'
        zhuisu_path_tmp = curpath / 'document_templates' / '报告追踪模板.docx'
        if zhuisu_path_tmp.is_file():
            self.sin_out.emit('已检测到有报告追溯模板文件...')
        else:
            self.sin_out.emit('open failed:选择的文档')
            return

        # 创建个列表放数据
        data_list = []
        # 由于docx的是列表，所以直接len函数统计count
        count = len(doc.tables)
        self.sin_out.emit('total:' + str(count))
        k = 0
        for tb in doc.tables:
            k += 1
            self.sin_out.emit('total:' + str(k))
            #注意docx处理方式不一样从0开始，并且要算总行数
            try:
                if tb.cell(1,1).text.find('测试用例名称') != -1:
                    data = {'yongli_ming':'','yongli_biaoshi':'','yongli_qingkuang':'','beizhu':''}
                    data['yongli_ming'] = tb.cell(1,4).text
                    data['yongli_biaoshi'] = tb.cell(1,8).text
                    wenti = tb.rows[-2].cells[2]
                    print('提取出来的信息：',wenti.text)
                    if wenti.text == '/' or wenti.text == '':
                        data['yongli_qingkuang'] = '通过'
                        data['beizhu'] = '/'
                    else:
                        data['yongli_qingkuang'] = '不通过'
                        data['beizhu'] = wenti.text
                    self.sin_out.emit(f'处理完毕({tb.cell(1,8).text})用例..')
                    data_list.append(data)
                else:
                    self.sin_out.emit(f'当前表格({tb.cell(1,8).text})用例无法识别请检查')
            except:
                self.sin_out.emit(f'处理第{k}个表格失败，请检查该表格...')
                pass

        try:
            tpl_path = Path.cwd() / "need" / "document_templates" / '报告追踪模板.docx'
            self.sin_out.emit('导入模板文件路径为：' + str(tpl_path))
            tpl = DocxTemplate(tpl_path)  # 模板导入成功

        except:
            QMessageBox.warning(self.parent, "出错了", "导入模板出错请检查模板文件是否存在或名字不正确")
            return

        # 开始渲染模板文件
        print(data_list)
        try:
            context = {
                "tables": data_list,
            }
            tpl.render(context)
            tpl.save("说明追踪文档.docx")
            QMessageBox.about(self.parent, "生成文档成功", "请查看当前工具根目录（报告追踪文档.docx）")
            self.sin_out.emit('stopthread')
        except:
            QMessageBox.warning(self.parent, "生成文档出错", "生成文档错误，请确认模板文档是否已打开或格式错误")
            self.sin_out.emit('stopthread')
            return

##################################################################################
##
##################################################################################