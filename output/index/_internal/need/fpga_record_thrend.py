import pythoncom
from PyQt5 import QtCore
from PyQt5.QtCore import pyqtSignal
from pathlib import *
from docx import Document
import re
from docx.shared import Pt

#常量
TABLE_FONT_SIZE = Pt(10.5)

class create_FPGA_record(QtCore.QThread):
    sin_out = pyqtSignal(str)

    def __init__(self,parent):
        super().__init__()
        self.parent = parent
        
    def run(self):
        
        
        self.sin_out.emit("开始填写FPGA记录......")
        self.sin_out.emit("开始填写文档......")
         
        #如果没有选择路径则退出
        if not self.parent.open_file_name:
            self.sin_out.emit('nofile')
            self.parent.tabWidget.setEnabled(True)
            return
        
        #告诉windows单线程
        pythoncom.CoInitialize()
        #在用户选择的目录中查找UAS单位测试报告文档
        self.sin_out.emit('打开单元测试原文件...')
        
        try:
            doc = Document(self.parent.open_file_name[0])
        except:
            self.sin_out.emit('open failed:选择的文档')
            pythoncom.CoUninitialize()
            self.parent.tabWidget.setEnabled(True)
            return
        
        table_count = len(doc.tables)
        self.sin_out.emit('total:'+ str(table_count))
        # 如果大纲标识一致需要累加
        static_dagang_biaoshi = ''
        index = 1
        num = 0
        for table in doc.tables:
            num += 1
            self.sin_out.emit(str(num))
            self.sin_out.emit(f'正在处理第{num}个表格')
            if table.cell(0,0).text == '测试用例名称':
                
                try:
                    #~~~~~~~~~~~~~~~第一步处理表格中所在标题填入~~~~~~~~~~~~~~~
                    prev_para = table._element.getprevious()
                    # 如果找到前一个元素是paragraph、但没有找到子节点有pStyle节点，则再往上找
                    while prev_para is not None and prev_para.tag.endswith('}p') and \
                        len(prev_para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')) == 0:
                            prev_para = prev_para.getprevious()
                    if prev_para is not None and prev_para.tag.endswith('}p'):
                        t_ele = prev_para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        title = '' 
                        for i in range(len(t_ele)):
                            title = title + t_ele[i].text
                        # 将title放入到cell(0,1)-完成第一步了
                        table.cell(0,3).text = title
                        
                        
                    #~~~~~~~~~~~~~~~第二步找到追踪关系中大纲标识更改后填入~~~~~~~~~~~~~~~(要求用户必须3行！)
                    temp = table.cell(1,3).text.split("\n")[-1]
                    dagang_biaoshi = re.split('[:：]',temp)[-1]
                    
                    # 如果大纲标识一致，则累加
                    if dagang_biaoshi != static_dagang_biaoshi:
                        static_dagang_biaoshi = dagang_biaoshi
                        yongli_biaoshi = dagang_biaoshi.replace('XQ','R1_YL') + '_001'
                        # 填入标识
                        table.cell(0,8).text = yongli_biaoshi
                        index = 1
                    else:
                        index += 1
                        if len(str(index)) <= 3:
                            str_index = (3-len(str(index)))*'0' + str(index)
                        yongli_biaoshi = dagang_biaoshi.replace('XQ','R1_YL') + '_' + str_index
                        table.cell(0,8).text = yongli_biaoshi
                    
                    #~~~~~~~~~~~~~~~输出用户知道~~~~~~~~~~~~~~~
                    self.sin_out.emit(f'处理完毕测试项：{title},{yongli_biaoshi}')
                        
                    #~~~~~~~~~~~~~~~第三步填写“测试用例综述”~~~~~~~~~~~~~~~
                    if title:
                        zongsu_string = f'使用功能仿真的方法，对{title}进行测试'
                        table.cell(2,3).text = zongsu_string
                    else:
                        print('注意：未识别到正确标题！，填入综述失败')
                        
                    #~~~~~~~~~~~~~~~第四步“执行状态”填写为已执行、“测试时间”、“测试人员”、“监测人员”为可更改项~~~~~~~~~~~~~~~
                    #执行状态-固定！
                    table.rows[-4].cells[2].text = '已执行'
                    #测试人员
                    table.rows[-3].cells[2].text = self.parent.lineEdit_8.text()
                    table.rows[-4].cells[7].text = self.parent.lineEdit_7.text()
                    table.rows[-3].cells[7].text = self.parent.lineEdit_15.text()
                    
                    #~~~~~~~~~~~~~~~第五步填写当通过与否为通过时，写"/"~~~~~~~~~~~~~~~
                    table.rows[-2].cells[2].text = '/'
                    
                    # 判断有几个步骤-根据总行数判断,并循环查看是否为通过
                    flag = 1
                    for i in range(len(table.rows) - 12):
                        temp = table.rows[-6 - i].cells[-1].text
                        if temp == '不通过':
                            flag = 0
                        if temp == '未通过':
                            flag = 0
                    if flag == 1:
                        table.rows[-2].cells[2].text = '/'
                    else:
                        project_code = self.parent.lineEdit_16.text()
                        pro_biaoshi = f'PT_{project_code}_'
                        table.rows[-2].cells[2].text = pro_biaoshi
                        
                    #~~~~~~~~~~~~~~~第六步截图为空填写‘/’否则不变~~~~~~~~~~~~~~~
                    cell = table.cell(-5,0)
                    has_image = False
                    for paragraph in cell.paragraphs:
                        tupian_list = paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                        if len(tupian_list) != 0:
                            has_image = True
                    if has_image:
                        pass
                    else:
                        table.cell(-5,0).text = '/'  
                except:
                    self.sin_out.emit(f'第{num}个表格处理失败，请查看！！！')
                    pass     
            
            # 设置字体
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = TABLE_FONT_SIZE
        # 保存文档
        try:
            doc.save('~新生产的fpga记录~.docx')
            self.sin_out.emit('stopsuccess')
        except:
            self.sin_out.emit('stoperror')
            return
        
        