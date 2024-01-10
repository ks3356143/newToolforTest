import tkinter
from tkinter import filedialog
from pathlib import Path
import re
from docxtpl import DocxTemplate
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# 匹配测试方法里面的标题
re_title = re.compile(r'（\w{2}_\w+_[a-zA-Z0-9]+_\w+\d+）')

window = tkinter.Tk()
window.title("大纲转换软件V0.01-tk-丑版")

# 窗口大小
window.geometry('600x300')

# label
label1 = tkinter.Label(window, text = '')
label1.grid(row = 0, column = 2)

def btn1click():
    path = tkinter.filedialog.askopenfilename()
    label1.config(text = path)

# 选择文件按钮
btn1 = tkinter.Button(window, text = "选择文件", command = btn1click)
btn1.grid(column = 1, row = 0)

# 生成文档函数
def generate_document():
    input_document_path = label1['text']
    # 定义当前四级标题的编号和名称
    current_level4_title = ''
    # 定义储存标题的东西
    level1_title = []
    level2_title = []
    level3_title = []
    level4_title = []
    level5_title = []
    level6_title = []
    level7_title = []
    if input_document_path:
        tpl_path = Path.cwd() / 'document_template' / '测试大纲生成测试说明模版.docx'
        tpl = DocxTemplate(tpl_path)  # 模板导入成功
        try:
            doc = Document(input_document_path)
        except:
            return
        shuoming_path_tmp = Path.cwd() / 'generate_document' / '生成的测试说明.docx'
        # 获取表格数量
        tables = doc.tables
        # 创建一个字典来储存单个用例
        data_list = []
        # 定义开关-当识别到测试定义的标题时候
        open_title = False
        for ele in doc._element.body.iter():
            if ele.tag.endswith('}p'):
                elePstyle = ele.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if len(elePstyle) >= 1:
                    parag = Paragraph(ele, doc)
                    if parag.style.name.startswith("Heading") or parag.style.name.startswith("标题"):
                        data = {}
                        rank = parag.style.name.split(" ")[-1]  # 标题等级str类型
                        text = parag.text  # 标题文字
                        # 先将标题储存在一个地方
                        if rank == '1' or rank == '10':
                            level1_title.append(1)
                            level2_title.clear()
                            level3_title.clear()
                            level4_title.clear()
                            level5_title.clear()
                            level6_title.clear()
                            level7_title.clear()
                        elif rank == '2' or rank == '20':
                            level2_title.append(1)
                            level3_title.clear()
                            level4_title.clear()
                            level5_title.clear()
                            level6_title.clear()
                            level7_title.clear()
                        elif rank == '3' or rank == '30':
                            level3_title.append(1)
                            level4_title.clear()
                            level5_title.clear()
                            level6_title.clear()
                            level7_title.clear()
                        elif rank == '4' or rank == '40':
                            level4_title.append(1)
                            level5_title.clear()
                            level6_title.clear()
                            level7_title.clear()
                            current_level4_title = f"({len(level1_title)}.{len(level2_title)}.{len(level3_title)}.{len(level4_title)}){text}"
                        elif rank == '5' or rank == '50':
                            level5_title.append(1)
                            level6_title.clear()
                            level7_title.clear()
                        elif rank == '6' or rank == '60':
                            level6_title.append(1)
                            level7_title.clear()
                        elif rank == '7' or rank == '70':
                            level7_title.append(1)
                        # 如果识别到标题为"测试定义"则打开获取标题开关
                        if text == '功能测试':
                            open_title = True
                        if text == '测试进度':
                            open_title = False
                        if open_title:
                            data['type'] = rank
                            data['text'] = text
                            data_list.append(data)
            elif ele.tag.endswith('}tbl'):
                table = Table(ele, doc)
                # 1识别到大纲的一个表格
                temp = table.cell(0, 0).text
                if temp == '测试项名称' and table.cell(0, 1).text != '文档审查' \
                    and table.cell(0, 1).text != '静态分析' and \
                        table.cell(0, 1).text != '代码审查':
                    # 先提取表格其他信息
                    csx_name = table.cell(0, 1).text
                    csx_ident = table.cell(0, 3).text
                    res_text = ''
                    prefix = ''
                    count = 1
                    for para in table.cell(4, 2).paragraphs:
                        # 2.1 识别到测试方法的标题
                        re_res = re_title.findall(para.text)
                        if re_res:
                            data = {}
                            count = 1
                            data['type'] = '5'
                            data['text'] = re.sub(r'\d{1,2}[.]', '', para.text).strip()
                            prefix = re_res[0].replace('（', '').replace("）", "")
                            res_text = data['text'].split('（')[0]
                            data_list.append(data)
                            # print('生成五级标题为:', data['text'])
                        # 2.2 如果不是标题则需要生成表格了
                        else:
                            tb_data = {'type': 'table'}
                            tb_data['name'] = f"{res_text}_{count}"
                            tb_data['ident'] = f"{prefix}_{count}"
                            tb_data['destination'] = f"验证{csx_name.replace('测试','')}是否正确"
                            tb_data['xqfx'] = current_level4_title
                            tb_data['xqident'] = prefix
                            tb_data['step'] = []
                            # 步骤处理共4个字段
                            # 要求每句话必须有；中文分号
                            split_temp = para.text.split('；')
                            if len(split_temp) == 1:
                                split_temp.append("")
                            index = 1
                            for sss in split_temp:
                                if sss:
                                    # 根据“查看分割”
                                    ck_split = sss.split('查看')
                                    if len(ck_split) == 1:
                                        ck_split.insert(0, "")
                                    step = {}
                                    step['index'] = index
                                    step['shuru'] = '巡天主控软件'
                                    step['guocheng'] = ck_split[0]
                                    # 去掉guocheng里面逗号
                                    if step['guocheng']:
                                        if step['guocheng'][-1] == '，':
                                            step['guocheng'] = step['guocheng'][:-1]
                                    else:
                                        print('请检查，未有查看字样...')
                                    step['yuqi'] = ck_split[1]
                                    tb_data['step'].append(step)
                                    index += 1

                            count += 1
                            data_list.append(tb_data)
        try:
            context = {
                "tables": data_list,
            }
            tpl.render(context)
            tpl.save(shuoming_path_tmp)
        except:
            print('在生成文档时报错')
    else:
        print('请先选择需要转换的文件')
    print('生成完成!')

# 生成文档按钮
btn2 = tkinter.Button(window, text = "生成文档", command = generate_document)
btn2.grid(column = 1, row = 1)

# 进入消息循环
window.mainloop()
